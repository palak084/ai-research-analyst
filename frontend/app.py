import streamlit as st
import requests
import PyPDF2
import docx
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import re
import io
import os
import urllib.parse
import time
import datetime
from fpdf import FPDF
from textblob import TextBlob
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation

API_BASE = os.environ.get("API_URL", "http://127.0.0.1:8000")

st.set_page_config(
    page_title="AI Research Analyst",
    page_icon="logo.png",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------- SESSION STATE INIT ----------
if "auth_passed" not in st.session_state:
    st.session_state.auth_passed = False
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "full_name" not in st.session_state:
    st.session_state.full_name = ""
if "token" not in st.session_state:
    st.session_state.token = ""
if "rag_chat_history" not in st.session_state:
    st.session_state.rag_chat_history = []
if "rag_doc_uploaded" not in st.session_state:
    st.session_state.rag_doc_uploaded = False
if "rag_doc_name" not in st.session_state:
    st.session_state.rag_doc_name = ""
if "last_analysis" not in st.session_state:
    st.session_state.last_analysis = None
if "last_analysis_text" not in st.session_state:
    st.session_state.last_analysis_text = ""
# current_mode: None = show landing, or "Research Analysis", "Chat with Document", "Compare Documents"
if "current_mode" not in st.session_state:
    st.session_state.current_mode = None
# viewing_history_id: None = not viewing, or int = viewing a past analysis
if "viewing_history_id" not in st.session_state:
    st.session_state.viewing_history_id = None

# ---------- GLOBAL STYLING ----------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&family=Playfair+Display:wght@600;700&display=swap');

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
    color: #3B2F24 !important;
}

p, span, label, li, ol, ul, h1, h2, h3, h4, h5, h6,
div, .stMarkdown, .stText,
[data-testid="stMarkdownContainer"],
[data-testid="stMarkdownContainer"] p,
[data-testid="stText"] {
    color: #3B2F24 !important;
}

#MainMenu {visibility: hidden;}
footer {visibility: hidden;}

.stApp { background: #FDF8F0; }

section[data-testid="stSidebar"] {
    background: #F3E4C9;
    border-right: 1px solid rgba(169, 139, 118, 0.3);
}
section[data-testid="stSidebar"] .block-container { padding-top: 2rem; }
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p {
    color: #5C4A3A !important;
}

.card {
    background: #FFFFFF;
    border: 1px solid rgba(169, 139, 118, 0.25);
    border-radius: 16px;
    padding: 1.5rem;
    margin-bottom: 1rem;
    box-shadow: 0 2px 12px rgba(169, 139, 118, 0.1);
    transition: border-color 0.3s ease, box-shadow 0.3s ease;
}
.card:hover {
    border-color: rgba(169, 139, 118, 0.5);
    box-shadow: 0 4px 20px rgba(169, 139, 118, 0.15);
}

.hero { text-align: center; padding: 2.5rem 0 0.5rem 0; }
.hero h1 {
    font-family: 'Playfair Display', serif;
    font-size: 2.6rem; font-weight: 700;
    background: linear-gradient(135deg, #5C4033, #8B6F5C, #4A3728);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    margin-bottom: 0.4rem;
}
.hero p { color: #5C4A3A !important; font-size: 1.05rem; font-weight: 500; }

.stat-row { display: flex; gap: 1rem; justify-content: center; flex-wrap: wrap; margin: 1.5rem 0; }
.stat-badge {
    background: #FFFFFF; border: 1px solid rgba(169, 139, 118, 0.25);
    border-radius: 14px; padding: 0.8rem 1.4rem; text-align: center;
    min-width: 150px; box-shadow: 0 2px 8px rgba(169, 139, 118, 0.08);
}
.stat-badge .num { font-size: 1.2rem; font-weight: 700; color: #3B2F24; }
.stat-badge .label { font-size: 0.7rem; color: #7B5E4B; text-transform: uppercase; letter-spacing: 1px; margin-top: 0.2rem; }

div.stButton > button {
    background: linear-gradient(135deg, #A98B76, #8B9B5A) !important;
    color: #FFFFFF !important; border: none !important;
    border-radius: 12px; height: 52px; width: 100%;
    font-size: 1.05rem; font-weight: 700;
    transition: all 0.3s ease;
    box-shadow: 0 4px 20px rgba(169, 139, 118, 0.25);
}
div.stButton > button:hover {
    background: linear-gradient(135deg, #8B7260, #6E7D45) !important;
    box-shadow: 0 6px 30px rgba(169, 139, 118, 0.4);
    transform: translateY(-1px); color: #FFFFFF !important;
}
div.stButton > button:active { transform: translateY(0); }
div.stButton > button p { color: #FFFFFF !important; font-weight: 700; }

.stTabs [data-baseweb="tab-list"] { gap: 0.5rem; background: rgba(169, 139, 118, 0.08); border-radius: 12px; padding: 0.3rem; }
.stTabs [data-baseweb="tab"] { border-radius: 10px; padding: 0.6rem 1.2rem; font-weight: 500; color: #7B5E4B !important; }
.stTabs [data-baseweb="tab"] p, .stTabs [data-baseweb="tab"] span { color: inherit !important; }
.stTabs [aria-selected="true"] { background: #FFFFFF !important; color: #3B2F24 !important; box-shadow: 0 2px 8px rgba(169, 139, 118, 0.15); }
.stTabs [aria-selected="true"] p, .stTabs [aria-selected="true"] span { color: #3B2F24 !important; }
.stTabs [data-baseweb="tab-highlight"] { display: none; }
.stTabs [data-baseweb="tab-border"] { display: none; }

.stTextArea textarea {
    background: #FFFFFF !important; border: 1px solid rgba(169, 139, 118, 0.3) !important;
    border-radius: 12px; color: #3B2F24 !important; font-size: 0.95rem; padding: 1rem; caret-color: #3B2F24;
}
.stTextArea textarea::placeholder { color: #A98B76 !important; opacity: 0.7 !important; }
.stTextArea textarea:focus { border-color: #A98B76 !important; box-shadow: 0 0 0 2px rgba(169, 139, 118, 0.15) !important; }
.stTextArea label { color: #5C4A3A !important; }

[data-testid="stFileUploader"] { background: #FFFFFF; border: 2px dashed rgba(169, 139, 118, 0.35); border-radius: 16px; padding: 1.5rem; }
[data-testid="stFileUploader"] span, [data-testid="stFileUploader"] p,
[data-testid="stFileUploader"] small, [data-testid="stFileUploader"] label { color: #5C4A3A !important; }
[data-testid="stFileUploader"] button { color: #3B2F24 !important; border-color: rgba(169, 139, 118, 0.4) !important; }

.stRadio > div { gap: 0.3rem; }
.stRadio label, .stRadio label span, .stRadio label p { color: #3B2F24 !important; font-size: 0.95rem; }
.stRadio [data-testid="stMarkdownContainer"] p { color: #3B2F24 !important; }

input, select, [data-baseweb="input"] input { color: #3B2F24 !important; background: #FFFFFF !important; }

[data-testid="stChatInput"] textarea, [data-testid="stChatInput"] input { color: #3B2F24 !important; background: #FFFFFF !important; }
[data-testid="stChatInput"] { background: #FFFFFF !important; border-color: rgba(169, 139, 118, 0.3) !important; }
[data-testid="stChatMessage"] { background: #FFFFFF !important; border: 1px solid rgba(169, 139, 118, 0.15); border-radius: 12px; padding: 1rem; margin-bottom: 0.5rem; }
[data-testid="stChatMessage"] p { color: #3B2F24 !important; }

[data-testid="stAlert"] p, .stAlert p { color: #3B2F24 !important; }
[data-testid="stSpinner"] p, .stSpinner p { color: #7B5E4B !important; }

.sidebar-label { font-size: 0.7rem; text-transform: uppercase; letter-spacing: 1.5px; color: #7B5E4B; margin-bottom: 0.3rem; }
.sidebar-title { font-family: 'Playfair Display', serif; font-size: 1.2rem; font-weight: 600; color: #3B2F24; margin-bottom: 1.5rem; }

.section-header {
    font-family: 'Playfair Display', serif; font-size: 1.1rem; font-weight: 600;
    color: #6B7D3E; margin-bottom: 0.5rem; padding-bottom: 0.5rem;
    border-bottom: 1px solid rgba(169, 139, 118, 0.2);
}

.result-content {
    background: #FAF5ED; border-radius: 12px; padding: 1.2rem;
    color: #3B2F24; line-height: 1.8; max-height: 500px; overflow-y: auto; font-size: 0.92rem;
}
.result-content::-webkit-scrollbar { width: 6px; }
.result-content::-webkit-scrollbar-track { background: transparent; }
.result-content::-webkit-scrollbar-thumb { background: rgba(169, 139, 118, 0.3); border-radius: 3px; }

.divider { height: 1px; background: linear-gradient(90deg, transparent, rgba(169,139,118,0.3), transparent); margin: 1.5rem 0; }
.stSpinner > div > div { color: #7B5E4B !important; }
div[data-testid="stAlert"] { border-radius: 12px; border: none; }

div[data-testid="stDownloadButton"] > button {
    background: linear-gradient(135deg, #8B9B5A, #A98B76) !important;
    color: #FFFFFF !important; border: none !important; border-radius: 12px;
    height: 48px; width: 100%; font-size: 1rem; font-weight: 700;
    box-shadow: 0 4px 20px rgba(139, 155, 90, 0.25);
}
div[data-testid="stDownloadButton"] > button:hover { box-shadow: 0 6px 30px rgba(139, 155, 90, 0.4); transform: translateY(-1px); }
div[data-testid="stDownloadButton"] > button p { color: #FFFFFF !important; font-weight: 700; }

[data-testid="stExpander"] { background: #FFFFFF; border: 1px solid rgba(169, 139, 118, 0.2); border-radius: 12px; }
[data-testid="stExpander"] summary p, [data-testid="stExpander"] summary span { color: #7B5E4B !important; font-weight: 500; }

.activity-log { background: #FFFFFF; border: 1px solid rgba(169,139,118,0.2); border-radius: 14px; padding: 1.2rem; margin-bottom: 1rem; }
.log-step { display: flex; align-items: flex-start; gap: 0.7rem; padding: 0.5rem 0; border-bottom: 1px solid rgba(169,139,118,0.08); }
.log-step:last-child { border-bottom: none; }
.log-dot { width: 10px; height: 10px; border-radius: 50%; margin-top: 5px; flex-shrink: 0; }
.log-dot.completed { background: #6B7D3E; }
.log-dot.running { background: #D4A030; }
.log-dot.system { background: #A98B76; }
.log-agent { font-weight: 600; font-size: 0.82rem; color: #5C4A3A; min-width: 130px; }
.log-action { font-size: 0.82rem; color: #3B2F24; flex: 1; }
.log-time { font-size: 0.72rem; color: #A98B76; white-space: nowrap; }

.web-result { background: #FFFFFF; border: 1px solid rgba(169,139,118,0.2); border-radius: 12px; padding: 1rem 1.2rem; margin-bottom: 0.6rem; transition: border-color 0.2s; }
.web-result:hover { border-color: #6B7D3E; }
.web-result-title { font-weight: 600; color: #3B2F24; font-size: 0.95rem; margin-bottom: 0.3rem; }
.web-result-snippet { font-size: 0.82rem; color: #5C4A3A; line-height: 1.5; margin-bottom: 0.3rem; }
.web-result-source { font-size: 0.72rem; color: #A98B76; }

.stTextInput input { background: #FFFFFF !important; border: 1px solid rgba(169,139,118,0.3) !important; border-radius: 10px !important; color: #3B2F24 !important; padding: 0.6rem 0.8rem !important; }
.stTextInput input:focus { border-color: #A98B76 !important; box-shadow: 0 0 0 2px rgba(169,139,118,0.15) !important; }

.history-item { background: #FFFFFF; border: 1px solid rgba(169,139,118,0.2); border-radius: 12px; padding: 1rem 1.2rem; margin-bottom: 0.6rem; }
.history-item:hover { border-color: #A98B76; box-shadow: 0 2px 12px rgba(169,139,118,0.15); }

.timeline-event { display: flex; gap: 1rem; margin-bottom: 1rem; padding-left: 1rem; border-left: 3px solid #A98B76; }
.timeline-date { font-weight: 700; color: #6B7D3E; font-size: 0.9rem; min-width: 100px; }
.timeline-text { font-size: 0.88rem; color: #3B2F24; line-height: 1.5; }

/* Sidebar history buttons */
section[data-testid="stSidebar"] div.stButton > button {
    background: rgba(255,255,255,0.6) !important;
    color: #3B2F24 !important;
    border: 1px solid rgba(169,139,118,0.2) !important;
    border-radius: 10px;
    height: auto !important;
    padding: 0.5rem 0.8rem !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    text-align: left !important;
    box-shadow: none !important;
    transition: all 0.2s ease;
}
section[data-testid="stSidebar"] div.stButton > button:hover {
    background: #FFFFFF !important;
    border-color: #A98B76 !important;
    transform: none !important;
    box-shadow: 0 2px 8px rgba(169,139,118,0.15) !important;
}
section[data-testid="stSidebar"] div.stButton > button p {
    color: #3B2F24 !important;
    font-weight: 500 !important;
}

/* New Chat button stands out */
section[data-testid="stSidebar"] div.stButton:first-of-type > button {
    background: linear-gradient(135deg, #A98B76, #8B9B5A) !important;
    color: #FFFFFF !important;
    border: none !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
    height: 44px !important;
    padding: 0 !important;
}
section[data-testid="stSidebar"] div.stButton:first-of-type > button:hover {
    background: linear-gradient(135deg, #8B7260, #6E7D45) !important;
}
section[data-testid="stSidebar"] div.stButton:first-of-type > button p {
    color: #FFFFFF !important;
    font-weight: 700 !important;
}
</style>
""", unsafe_allow_html=True)


# ---------- FILE EXTRACTION ----------
def extract_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_docx(file):
    doc = docx.Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_excel(file):
    df = pd.read_excel(file)
    return df.to_string()


# ---------- CHART GENERATION ----------
def generate_chart(text):
    words = re.findall(r'\w+', text.lower())
    keywords = ["ai", "data", "cost", "privacy", "healthcare",
                "model", "research", "analysis", "risk", "performance"]
    counts = [words.count(k) for k in keywords]

    filtered = [(k, c) for k, c in zip(keywords, counts) if c > 0]
    if not filtered:
        filtered = list(zip(keywords[:5], counts[:5]))

    labels, values = zip(*filtered)

    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor('#FDF8F0')
    ax.set_facecolor('#FDF8F0')

    colors = ['#A98B76', '#8B9B5A', '#BFA28C', '#6B7D3E'] * 3
    bars = ax.barh(labels, values, color=colors[:len(labels)],
                   edgecolor='#A98B76', linewidth=0.5, height=0.6)

    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height() / 2,
                str(val), va='center', color='#3B2F24', fontweight='600', fontsize=11)

    ax.set_xlabel('Frequency', color='#7B5E4B', fontsize=11)
    ax.set_title('Key Concept Frequency', color='#3B2F24', fontsize=14,
                 fontweight='600', pad=15, fontfamily='serif')
    ax.tick_params(colors='#5C4A3A', labelsize=10)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color('#D4C4B0')
    ax.spines['left'].set_color('#D4C4B0')
    ax.xaxis.grid(True, color='#D4C4B0', linestyle='--', alpha=0.5)
    ax.set_axisbelow(True)
    plt.tight_layout()
    return fig


# ---------- SENTIMENT ANALYSIS ----------
def analyze_sentiment(text):
    blob = TextBlob(text)
    polarity = blob.sentiment.polarity
    subjectivity = blob.sentiment.subjectivity

    if polarity > 0.1:
        tone = "Positive"
        tone_color = "#6B7D3E"
    elif polarity < -0.1:
        tone = "Negative"
        tone_color = "#B85C4A"
    else:
        tone = "Neutral"
        tone_color = "#A98B76"

    return polarity, subjectivity, tone, tone_color


def render_sentiment_gauge(polarity, subjectivity, tone, tone_color):
    fig, axes = plt.subplots(1, 2, figsize=(10, 4))
    fig.patch.set_facecolor('#FDF8F0')

    for ax in axes:
        ax.set_facecolor('#FDF8F0')

    # Polarity gauge
    ax = axes[0]
    colors_bar = ['#B85C4A', '#D4A030', '#A98B76', '#8B9B5A', '#6B7D3E']
    bounds = np.linspace(-1, 1, 6)
    for i in range(5):
        ax.barh(0, bounds[i + 1] - bounds[i], left=bounds[i], height=0.4,
                color=colors_bar[i], alpha=0.4)
    ax.axvline(polarity, color='#3B2F24', linewidth=3, zorder=5)
    ax.scatter([polarity], [0], color='#3B2F24', s=120, zorder=6)
    ax.set_xlim(-1, 1)
    ax.set_ylim(-0.5, 0.5)
    ax.set_yticks([])
    ax.set_title(f'Polarity: {polarity:.2f} ({tone})', color='#3B2F24',
                 fontweight='600', fontsize=12, fontfamily='serif')
    ax.set_xticks([-1, -0.5, 0, 0.5, 1])
    ax.set_xticklabels(['Very Negative', 'Negative', 'Neutral', 'Positive', 'Very Positive'],
                       fontsize=7, color='#5C4A3A')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_color('#D4C4B0')

    # Subjectivity gauge
    ax2 = axes[1]
    sub_colors = ['#6B7D3E', '#8B9B5A', '#A98B76', '#D4A030', '#B85C4A']
    sub_bounds = np.linspace(0, 1, 6)
    for i in range(5):
        ax2.barh(0, sub_bounds[i + 1] - sub_bounds[i], left=sub_bounds[i], height=0.4,
                 color=sub_colors[i], alpha=0.4)
    ax2.axvline(subjectivity, color='#3B2F24', linewidth=3, zorder=5)
    ax2.scatter([subjectivity], [0], color='#3B2F24', s=120, zorder=6)
    ax2.set_xlim(0, 1)
    ax2.set_ylim(-0.5, 0.5)
    ax2.set_yticks([])
    ax2.set_title(f'Subjectivity: {subjectivity:.2f}', color='#3B2F24',
                  fontweight='600', fontsize=12, fontfamily='serif')
    ax2.set_xticks([0, 0.25, 0.5, 0.75, 1])
    ax2.set_xticklabels(['Objective', '', 'Balanced', '', 'Subjective'],
                        fontsize=7, color='#5C4A3A')
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['left'].set_visible(False)
    ax2.spines['bottom'].set_color('#D4C4B0')

    plt.tight_layout()
    return fig


# ---------- TOPIC MODELING ----------
def extract_topics(text, n_topics=4, n_words=6):
    sentences = re.split(r'[.!?\n]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]

    if len(sentences) < n_topics:
        return None

    vectorizer = CountVectorizer(max_df=0.95, min_df=1, stop_words='english', max_features=500)
    try:
        dtm = vectorizer.fit_transform(sentences)
    except ValueError:
        return None

    lda = LatentDirichletAllocation(n_components=min(n_topics, len(sentences)),
                                    random_state=42, max_iter=10)
    lda.fit(dtm)

    feature_names = vectorizer.get_feature_names_out()
    topics = []
    for i, topic_weights in enumerate(lda.components_):
        top_indices = topic_weights.argsort()[-n_words:][::-1]
        top_words = [feature_names[j] for j in top_indices]
        weight = topic_weights[top_indices].sum()
        topics.append({"id": i + 1, "words": top_words, "weight": weight})

    return topics


def render_topic_chart(topics):
    if not topics:
        return None

    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor('#FDF8F0')
    ax.set_facecolor('#FDF8F0')

    labels = [f"Topic {t['id']}" for t in topics]
    weights = [t['weight'] for t in topics]
    colors = ['#A98B76', '#8B9B5A', '#6B7D3E', '#BFA28C', '#D4A030'] * 2
    word_labels = [", ".join(t['words'][:4]) for t in topics]

    bars = ax.barh(labels, weights, color=colors[:len(labels)],
                   edgecolor='#A98B76', linewidth=0.5, height=0.5)

    for bar, wl in zip(bars, word_labels):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height() / 2,
                wl, va='center', color='#5C4A3A', fontsize=9, style='italic')

    ax.set_xlabel('Topic Weight', color='#7B5E4B', fontsize=11)
    ax.set_title('Discovered Topics (LDA)', color='#3B2F24', fontsize=14,
                 fontweight='600', pad=15, fontfamily='serif')
    ax.tick_params(colors='#5C4A3A', labelsize=10)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color('#D4C4B0')
    ax.spines['left'].set_color('#D4C4B0')
    ax.xaxis.grid(True, color='#D4C4B0', linestyle='--', alpha=0.5)
    ax.set_axisbelow(True)
    plt.tight_layout()
    return fig


# ---------- RESEARCH TIMELINE ----------
def extract_timeline(text):
    events = []

    # Match patterns like: "In 2020, ...", "By March 2021, ...", "2019: ...", "(2023)"
    patterns = [
        (r'(?:In|Since|By|During|From|After|Before|Around)\s+'
         r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+)?'
         r'(\d{4})\s*[,.]?\s*(.{15,120}?)(?:\.|$)', 'full'),
        (r'(\d{4})\s*[-:]\s*(.{15,120}?)(?:\.|$)', 'year_colon'),
        (r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})\s*[,.:]\s*(.{15,120}?)(?:\.|$)', 'full_date'),
    ]

    seen_years = set()
    for pattern, ptype in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for m in matches:
            if ptype == 'full':
                month = m.group(1) or ""
                year = m.group(2)
                event_text = m.group(3).strip()
                date_str = f"{month}{year}".strip()
            elif ptype == 'year_colon':
                year = m.group(1)
                event_text = m.group(2).strip()
                date_str = year
            elif ptype == 'full_date':
                date_str = m.group(1)
                year = re.search(r'\d{4}', date_str).group()
                event_text = m.group(2).strip()

            if year not in seen_years and len(event_text) > 15:
                events.append({"date": date_str, "year": int(year), "text": event_text})
                seen_years.add(year)

    events.sort(key=lambda x: x["year"])
    return events[:12]


# ---------- CITATION EXTRACTION ----------
def extract_citations(text):
    citations = []

    numbered = re.findall(r'\[(\d+)\]\s*(.+?)(?:\n|$)', text)
    for num, ref in numbered:
        ref = ref.strip().rstrip('.')
        citations.append({"id": f"[{num}]", "text": ref})

    apa = re.findall(
        r'([A-Z][a-z]+(?:,?\s+(?:[A-Z]\.?\s*)+)?(?:,?\s*(?:&|and)\s+[A-Z][a-z]+(?:,?\s+(?:[A-Z]\.?\s*)+)?)*)\s*\((\d{4})\)\.\s*(.+?)(?:\.\s|$)',
        text
    )
    for author, year, title in apa:
        ref_text = f"{author} ({year}). {title.strip().rstrip('.')}"
        if not any(ref_text in c["text"] for c in citations):
            citations.append({"id": f"({year})", "text": ref_text})

    inline = re.findall(
        r'\(([A-Z][a-z]+(?:\s+(?:et\s+al\.))?(?:\s*(?:&|and)\s*[A-Z][a-z]+)*),?\s*(\d{4})\)',
        text
    )
    for author, year in inline:
        ref_text = f"{author}, {year}"
        if not any(ref_text in c["text"] for c in citations):
            citations.append({"id": f"({year})", "text": ref_text})

    ref_section = re.search(
        r'(?:References|Bibliography|Works Cited)\s*\n([\s\S]+?)(?:\n\n|\Z)',
        text, re.IGNORECASE
    )
    if ref_section:
        for line in ref_section.group(1).strip().split('\n'):
            line = line.strip().lstrip('\u2022-\u2013\u2014 ')
            if len(line) > 15 and not any(line in c["text"] for c in citations):
                citations.append({"id": "ref", "text": line.rstrip('.')})

    for c in citations:
        query = urllib.parse.quote_plus(c["text"][:120])
        c["scholar_url"] = f"https://scholar.google.com/scholar?q={query}"

    return citations


# ---------- PDF REPORT GENERATION ----------
class ResearchPDF(FPDF):
    def header(self):
        self.set_fill_color(44, 36, 24)
        self.rect(0, 0, 210, 30, 'F')
        self.set_font('Helvetica', 'B', 18)
        self.set_text_color(243, 228, 201)
        self.cell(0, 15, 'AI Research Analyst', align='C', new_x="LMARGIN", new_y="NEXT")
        self.set_font('Helvetica', '', 9)
        self.set_text_color(169, 139, 118)
        self.cell(0, 6, 'Multi-agent AI Research Report', align='C', new_x="LMARGIN", new_y="NEXT")
        self.ln(8)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(169, 139, 118)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')

    def section_title(self, title):
        self.set_font('Helvetica', 'B', 13)
        self.set_text_color(107, 125, 62)
        self.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(169, 139, 118)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(3)

    def section_body(self, body):
        self.set_font('Helvetica', '', 10)
        self.set_text_color(60, 50, 40)
        self.multi_cell(0, 5.5, body)
        self.ln(4)


def generate_pdf(plan, analysis, insights, chart_fig, citations=None):
    pdf = ResearchPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    pdf.section_title('Research Plan')
    pdf.section_body(plan)
    pdf.section_title('Detailed Analysis')
    pdf.section_body(analysis)
    pdf.section_title('Key Insights')
    pdf.section_body(insights)

    if chart_fig:
        pdf.section_title('Concept Frequency Chart')
        buf = io.BytesIO()
        chart_fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#FDF8F0')
        buf.seek(0)
        pdf.image(buf, x=15, w=180)
        pdf.ln(6)

    if citations:
        pdf.section_title('Extracted Citations')
        pdf.set_font('Helvetica', '', 9)
        for c in citations:
            pdf.set_text_color(60, 50, 40)
            pdf.multi_cell(0, 5, f"{c['id']}  {c['text']}")
            pdf.set_text_color(100, 120, 160)
            pdf.set_font('Helvetica', 'U', 8)
            pdf.multi_cell(0, 4, c['scholar_url'])
            pdf.set_font('Helvetica', '', 9)
            pdf.ln(2)

    return bytes(pdf.output())


# ---------- ACTIVITY LOG RENDERER ----------
def render_activity_log(activity_log):
    if not activity_log:
        return
    start_time = activity_log[0]["timestamp"]
    html = '<div class="activity-log"><div class="section-header">Agent Activity Log</div>'
    for step in activity_log:
        elapsed = step["timestamp"] - start_time
        agent = step["agent"]
        action = step["action"]
        step_status = step.get("status", "completed")
        dot_class = "system" if agent == "System" else "completed"
        if step_status == "running":
            dot_class = "running"
        html += f'<div class="log-step"><div class="log-dot {dot_class}"></div>'
        html += f'<div class="log-agent">{agent}</div>'
        html += f'<div class="log-action">{action}</div>'
        html += f'<div class="log-time">+{elapsed:.1f}s</div></div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)


# ---------- RESULTS RENDERER ----------
def render_results(data, text_input, show_save=False):
    citations = extract_citations(text_input)
    analysis_citations = extract_citations(data.get("analysis", "") + "\n" + data.get("insights", ""))
    for c in analysis_citations:
        if not any(c["text"] in existing["text"] for existing in citations):
            citations.append(c)

    if "activity_log" in data and data["activity_log"]:
        render_activity_log(data["activity_log"])

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # Prepare all analysis data
    combined_text = text_input + "\n" + data.get("analysis", "") + "\n" + data.get("insights", "")
    polarity, subjectivity, tone, tone_color = analyze_sentiment(combined_text)
    topics = extract_topics(combined_text)
    timeline_events = extract_timeline(text_input)
    web_results = data.get("web_results", [])

    tab_names = ["Research Plan", "Analysis", "Insights", "Visualizations",
                 "Sentiment", "Topics", "Timeline", "Citations"]
    if web_results:
        tab_names.append("Web Results")

    tabs = st.tabs(tab_names)

    with tabs[0]:
        st.markdown(f'<div class="card"><div class="section-header">Research Plan</div>'
                    f'<div class="result-content">{data["plan"]}</div></div>', unsafe_allow_html=True)

    with tabs[1]:
        st.markdown(f'<div class="card"><div class="section-header">Detailed Analysis</div>'
                    f'<div class="result-content">{data["analysis"]}</div></div>', unsafe_allow_html=True)

    with tabs[2]:
        st.markdown(f'<div class="card"><div class="section-header">Key Insights</div>'
                    f'<div class="result-content">{data["insights"]}</div></div>', unsafe_allow_html=True)

    with tabs[3]:
        st.markdown('<div class="card"><div class="section-header">Concept Frequency</div></div>',
                    unsafe_allow_html=True)
        fig = generate_chart(data["analysis"])
        st.pyplot(fig)

    # Sentiment tab
    with tabs[4]:
        st.markdown(f'<div class="card"><div class="section-header">Sentiment &amp; Tone Analysis</div>'
                    f'<div style="display:flex; gap:2rem; flex-wrap:wrap; margin:1rem 0;">'
                    f'<div style="text-align:center;">'
                    f'<div style="font-size:2rem; font-weight:700; color:{tone_color};">{tone}</div>'
                    f'<div style="font-size:0.8rem; color:#7B5E4B;">Overall Tone</div></div>'
                    f'<div style="text-align:center;">'
                    f'<div style="font-size:1.5rem; font-weight:700; color:#3B2F24;">{polarity:+.2f}</div>'
                    f'<div style="font-size:0.8rem; color:#7B5E4B;">Polarity (-1 to +1)</div></div>'
                    f'<div style="text-align:center;">'
                    f'<div style="font-size:1.5rem; font-weight:700; color:#3B2F24;">{subjectivity:.2f}</div>'
                    f'<div style="font-size:0.8rem; color:#7B5E4B;">Subjectivity (0 to 1)</div></div>'
                    f'</div></div>', unsafe_allow_html=True)
        sentiment_fig = render_sentiment_gauge(polarity, subjectivity, tone, tone_color)
        st.pyplot(sentiment_fig)

    # Topics tab
    with tabs[5]:
        if topics:
            st.markdown('<div class="card"><div class="section-header">Topic Modeling (LDA)</div></div>',
                        unsafe_allow_html=True)
            topic_fig = render_topic_chart(topics)
            if topic_fig:
                st.pyplot(topic_fig)

            for t in topics:
                words_html = " ".join(
                    [f'<span style="background:#FAF5ED; border:1px solid rgba(169,139,118,0.2); '
                     f'border-radius:8px; padding:0.2rem 0.6rem; font-size:0.82rem; color:#3B2F24; '
                     f'margin:0.2rem;">{w}</span>' for w in t['words']]
                )
                st.markdown(
                    f'<div class="card" style="padding:1rem 1.2rem;">'
                    f'<div style="font-weight:600; color:#6B7D3E; margin-bottom:0.4rem;">Topic {t["id"]}</div>'
                    f'<div style="display:flex; flex-wrap:wrap; gap:0.3rem;">{words_html}</div></div>',
                    unsafe_allow_html=True
                )
        else:
            st.markdown(
                '<div class="card" style="text-align:center; padding:2rem;">'
                '<div style="color:#7B5E4B;">Not enough text to extract topics. '
                'Try with longer documents.</div></div>', unsafe_allow_html=True
            )

    # Timeline tab
    with tabs[6]:
        if timeline_events:
            st.markdown(f'<div class="card"><div class="section-header">'
                        f'{len(timeline_events)} Events Found</div></div>', unsafe_allow_html=True)
            for ev in timeline_events:
                st.markdown(
                    f'<div class="timeline-event">'
                    f'<div class="timeline-date">{ev["date"]}</div>'
                    f'<div class="timeline-text">{ev["text"]}</div></div>',
                    unsafe_allow_html=True
                )
        else:
            st.markdown(
                '<div class="card" style="text-align:center; padding:2rem;">'
                '<div style="color:#7B5E4B;">No datable events found in the text.<br>'
                '<span style="font-size:0.8rem; opacity:0.7;">'
                'Tip: Include dates like "In 2020, ..." or "March 2021: ..."</span></div></div>',
                unsafe_allow_html=True
            )

    # Citations tab
    with tabs[7]:
        if citations:
            st.markdown(f'<div class="card"><div class="section-header">'
                        f'{len(citations)} Citation(s) Found</div></div>', unsafe_allow_html=True)
            for c in citations:
                st.markdown(
                    f'<div class="card" style="padding:1rem 1.2rem;">'
                    f'<div style="color:#3B2F24; font-size:0.92rem; line-height:1.6;">'
                    f'<span style="color:#6B7D3E; font-weight:600;">{c["id"]}</span>'
                    f'&nbsp; {c["text"]}</div>'
                    f'<a href="{c["scholar_url"]}" target="_blank"'
                    f' style="color:#6B7D3E; font-size:0.8rem; text-decoration:none; opacity:0.8;">'
                    f'Search on Google Scholar &rarr;</a></div>',
                    unsafe_allow_html=True
                )
        else:
            st.markdown(
                '<div class="card" style="text-align:center; padding:2rem;">'
                '<div style="color:#7B5E4B; font-size:0.95rem;">'
                'No citations detected in the input text.<br>'
                '<span style="font-size:0.8rem; opacity:0.7;">'
                'Tip: Include a References section or use formats like '
                '(Author, Year) or [1] Author, Title...</span></div></div>',
                unsafe_allow_html=True
            )

    # Web results tab
    if web_results and len(tabs) > 8:
        with tabs[8]:
            st.markdown(f'<div class="card"><div class="section-header">'
                        f'{len(web_results)} Related Papers &amp; Articles</div></div>',
                        unsafe_allow_html=True)
            for wr in web_results:
                st.markdown(
                    f'<div class="web-result">'
                    f'<a href="{wr.get("url", "#")}" target="_blank" style="text-decoration:none;">'
                    f'<div class="web-result-title">{wr.get("title", "Untitled")}</div></a>'
                    f'<div class="web-result-snippet">{wr.get("snippet", "")}</div>'
                    f'<div class="web-result-source">{wr.get("source", "")}</div></div>',
                    unsafe_allow_html=True
                )

    # PDF Download
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    fig_for_pdf = generate_chart(data["analysis"])
    pdf_bytes = generate_pdf(
        data["plan"], data["analysis"], data["insights"],
        fig_for_pdf, citations if citations else None
    )
    st.download_button(
        label="Download PDF Report", data=pdf_bytes,
        file_name="research_report.pdf", mime="application/pdf",
    )

    # Auto-saved indicator
    if st.session_state.logged_in and not show_save:
        pass  # viewing history, no indicator needed
    elif st.session_state.logged_in:
        st.markdown(
            '<div style="text-align:center; color:#6B7D3E; font-size:0.8rem; padding:0.5rem;">'
            '&#10003; Auto-saved to your history</div>',
            unsafe_allow_html=True
        )


# ====================================================================
# AUTH GATE — full-page login / register / continue as guest
# ====================================================================
if not st.session_state.auth_passed:
    st.markdown("""
    <div style="text-align:center; padding:2.5rem 0 1rem 0;">
        <div style="font-family:'Playfair Display',serif; font-size:2.4rem; font-weight:700;
                    background:linear-gradient(135deg,#7B5E4B,#A98B76,#5C4A3A);
                    -webkit-background-clip:text; -webkit-text-fill-color:transparent;
                    margin-bottom:0.3rem;">
            AI Research Analyst
        </div>
        <div style="color:#7B5E4B; font-size:0.95rem; margin-bottom:0.5rem;">
            Multi-agent AI system for research analysis
        </div>
    </div>
    """, unsafe_allow_html=True)

    _, center_col, _ = st.columns([1.2, 1.6, 1.2])
    with center_col:
        auth_choice = st.radio(
            "Choose how to continue",
            ["Sign In", "Create Account", "Continue as Guest"],
            label_visibility="collapsed",
            horizontal=True
        )

        if auth_choice == "Sign In":
            login_user = st.text_input("Username", key="login_user", placeholder="Enter your username")
            login_pass = st.text_input("Password", type="password", key="login_pass", placeholder="Enter your password")
            if st.button("Sign In", key="login_btn", use_container_width=True):
                if login_user and login_pass:
                    try:
                        resp = requests.post(f"{API_BASE}/auth/login", json={
                            "username": login_user, "password": login_pass
                        })
                        if resp.status_code == 200:
                            data = resp.json()
                            st.session_state.logged_in = True
                            st.session_state.auth_passed = True
                            st.session_state.username = data["user"]["username"]
                            st.session_state.full_name = data["user"]["full_name"]
                            st.session_state.token = data["token"]
                            st.rerun()
                        else:
                            st.error("Invalid username or password")
                    except Exception:
                        st.error("Cannot connect to server. Is the backend running?")
                else:
                    st.warning("Please fill in both fields")

        elif auth_choice == "Create Account":
            reg_name = st.text_input("Full Name", key="reg_name", placeholder="Your full name")
            reg_user = st.text_input("Username", key="reg_user", placeholder="Choose a username")
            reg_pass = st.text_input("Password", type="password", key="reg_pass", placeholder="Choose a password")
            if st.button("Create Account", key="reg_btn", use_container_width=True):
                if reg_user and reg_pass:
                    try:
                        resp = requests.post(f"{API_BASE}/auth/register", json={
                            "username": reg_user, "password": reg_pass, "full_name": reg_name
                        })
                        if resp.status_code == 200:
                            data = resp.json()
                            st.session_state.logged_in = True
                            st.session_state.auth_passed = True
                            st.session_state.username = data["user"]["username"]
                            st.session_state.full_name = data["user"]["full_name"]
                            st.session_state.token = data["token"]
                            st.rerun()
                        else:
                            st.error(resp.json().get("detail", "Registration failed"))
                    except Exception:
                        st.error("Cannot connect to server. Is the backend running?")
                else:
                    st.warning("Please fill in username and password")

        else:  # Continue as Guest
            st.markdown(
                '<div style="color:#7B5E4B; font-size:0.85rem; margin:0.5rem 0 1rem 0;">'
                'You can analyze documents without an account. Saving analyses to history requires login.</div>',
                unsafe_allow_html=True
            )
            if st.button("Continue as Guest", key="guest_btn", use_container_width=True):
                st.session_state.auth_passed = True
                st.session_state.logged_in = False
                st.rerun()

    st.stop()


# ---------- SIDEBAR (shown only after auth) ----------
with st.sidebar:
    # User info
    if st.session_state.logged_in:
        st.markdown(
            f'<div style="background:#FFFFFF; border-radius:10px; padding:0.8rem 1rem;'
            f' border:1px solid rgba(169,139,118,0.2); margin-bottom:0.8rem;">'
            f'<div style="font-weight:600; color:#3B2F24; font-size:0.9rem;">'
            f'{st.session_state.full_name or st.session_state.username}</div>'
            f'<div style="font-size:0.75rem; color:#7B5E4B;">@{st.session_state.username}</div></div>',
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            '<div style="background:#FFFFFF; border-radius:10px; padding:0.6rem 1rem;'
            ' border:1px solid rgba(169,139,118,0.2); margin-bottom:0.8rem;'
            ' font-size:0.82rem; color:#7B5E4B;">Guest Mode</div>',
            unsafe_allow_html=True
        )

    # New Chat button
    if st.button("New Chat", key="new_chat_btn", use_container_width=True):
        st.session_state.current_mode = None
        st.session_state.last_analysis = None
        st.session_state.last_analysis_text = ""
        st.session_state.viewing_history_id = None
        st.session_state.rag_chat_history = []
        st.session_state.rag_doc_uploaded = False
        st.session_state.rag_doc_name = ""
        st.rerun()

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # Past chats (history)
    if st.session_state.logged_in:
        st.markdown('<div class="sidebar-label">Recent</div>', unsafe_allow_html=True)
        try:
            hist_resp = requests.get(f"{API_BASE}/history/{st.session_state.username}")
            if hist_resp.status_code == 200:
                history_items = hist_resp.json()
                if history_items:
                    for item in history_items[:15]:
                        title_short = item["title"][:35] + ("..." if len(item["title"]) > 35 else "")
                        if st.button(
                            title_short,
                            key=f"sidebar_hist_{item['id']}",
                            use_container_width=True,
                        ):
                            st.session_state.viewing_history_id = item["id"]
                            st.session_state.current_mode = None
                            st.session_state.last_analysis = None
                            st.rerun()
                else:
                    st.markdown(
                        '<div style="color:#A98B76; font-size:0.8rem; padding:0.5rem 0;">'
                        'No past research yet</div>',
                        unsafe_allow_html=True
                    )
        except Exception:
            pass

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # Logout at bottom
    if st.session_state.logged_in:
        if st.button("Logout", key="logout_btn"):
            st.session_state.logged_in = False
            st.session_state.auth_passed = False
            st.session_state.username = ""
            st.session_state.full_name = ""
            st.session_state.token = ""
            st.rerun()

    st.markdown(
        '<div style="font-size: 0.75rem; color: #7B5E4B; opacity: 0.6;">'
        'Powered by Ollama &bull; FastAPI &bull; Streamlit</div>',
        unsafe_allow_html=True
    )

# Resolve current mode
mode = st.session_state.current_mode


# ====================================================================
# VIEWING A PAST CHAT FROM HISTORY
# ====================================================================
if st.session_state.viewing_history_id:
    try:
        detail_resp = requests.get(
            f"{API_BASE}/history/detail/{st.session_state.viewing_history_id}"
        )
        if detail_resp.status_code == 200:
            detail = detail_resp.json()
            dt = datetime.datetime.fromtimestamp(detail.get("created_at", 0))
            date_str = dt.strftime("%B %d, %Y at %I:%M %p")

            plan_text = detail.get("plan", "")
            is_chat = plan_text.startswith("Chat with Document:")
            is_compare = plan_text.startswith("Compare Documents:")

            st.markdown(
                f'<div class="hero"><h1>{detail["title"][:60]}</h1>'
                f'<p>Saved on {date_str}</p></div>',
                unsafe_allow_html=True
            )
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

            if is_chat:
                # Render as chat conversation
                source_label = plan_text.replace("Chat with Document: ", "")
                st.markdown(
                    f'<div style="display:flex; align-items:center; background:#FFFFFF;'
                    f' border:1px solid rgba(169,139,118,0.2); border-radius:12px;'
                    f' padding:0.7rem 1.2rem; margin-bottom:1rem;">'
                    f'<span style="color:#6B7D3E; font-size:1.1rem; margin-right:0.6rem;">&#128172;</span>'
                    f'<span style="color:#3B2F24; font-weight:600; font-size:0.9rem;">'
                    f'Chat with Document: {source_label}</span></div>',
                    unsafe_allow_html=True
                )

                # Show question
                st.markdown(
                    f'<div style="background:#FFFFFF; border:1px solid rgba(169,139,118,0.15);'
                    f' border-radius:12px; padding:1rem; margin-bottom:0.5rem;">'
                    f'<div style="color:#7B5E4B; font-size:0.75rem; margin-bottom:0.3rem;">You asked:</div>'
                    f'<div style="color:#3B2F24; font-size:0.95rem;">{detail.get("input_text", "")}</div></div>',
                    unsafe_allow_html=True
                )

                # Show answer
                analysis_text = detail.get("analysis", "")
                st.markdown(
                    f'<div style="background:#FAF5ED; border:1px solid rgba(169,139,118,0.15);'
                    f' border-radius:12px; padding:1rem; margin-bottom:0.5rem;">'
                    f'<div style="color:#6B7D3E; font-size:0.75rem; margin-bottom:0.3rem;">AI Response:</div>'
                    f'<div style="color:#3B2F24; font-size:0.95rem; line-height:1.7;">{analysis_text}</div></div>',
                    unsafe_allow_html=True
                )

            elif is_compare:
                # Render as comparison
                doc_names_str = plan_text.replace("Compare Documents: ", "")
                st.markdown(
                    f'<div style="display:flex; align-items:center; background:#FFFFFF;'
                    f' border:1px solid rgba(169,139,118,0.2); border-radius:12px;'
                    f' padding:0.7rem 1.2rem; margin-bottom:1rem;">'
                    f'<span style="color:#6B7D3E; font-size:1.1rem; margin-right:0.6rem;">&#128203;</span>'
                    f'<span style="color:#3B2F24; font-weight:600; font-size:0.9rem;">'
                    f'Compared: {doc_names_str}</span></div>',
                    unsafe_allow_html=True
                )

                tab1, tab2 = st.tabs(["Document Summaries", "Comparison"])
                with tab1:
                    st.markdown(
                        f'<div class="card"><div class="section-header">Summaries</div>'
                        f'<div class="result-content">{detail.get("analysis", "")}</div></div>',
                        unsafe_allow_html=True
                    )
                with tab2:
                    st.markdown(
                        f'<div class="card"><div class="section-header">Comparative Analysis</div>'
                        f'<div class="result-content">{detail.get("insights", "")}</div></div>',
                        unsafe_allow_html=True
                    )

            else:
                # Regular research analysis
                render_results(detail, detail.get("input_text", ""), show_save=False)

            # Delete option
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            if st.button("Delete This Analysis"):
                try:
                    requests.delete(f"{API_BASE}/history/{st.session_state.viewing_history_id}")
                    st.session_state.viewing_history_id = None
                    st.rerun()
                except Exception:
                    st.error("Failed to delete.")
        else:
            st.error("Could not load this analysis.")
    except Exception:
        st.error("Could not connect to backend.")


# ====================================================================
# LANDING PAGE — no mode selected (New Chat)
# ====================================================================
elif mode is None:
    st.markdown("""
    <div class="hero">
        <h1>AI Research Analyst</h1>
        <p>What would you like to do?</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("")

    col1, col2, col3 = st.columns(3, gap="medium")

    with col1:
        st.markdown(
            '<div class="card" style="text-align:center; padding:2rem; min-height:220px; cursor:pointer;">'
            '<div style="font-size:2.2rem; margin-bottom:0.6rem;">&#128300;</div>'
            '<div style="font-family:Playfair Display,serif; font-size:1.15rem; font-weight:600;'
            ' color:#3B2F24; margin-bottom:0.4rem;">Research Analysis</div>'
            '<div style="font-size:0.82rem; color:#7B5E4B; line-height:1.5;">'
            'Paste text or upload a document for multi-agent AI analysis with insights, '
            'charts, and PDF export</div></div>',
            unsafe_allow_html=True
        )
        if st.button("Start Analysis", key="land_analysis", use_container_width=True):
            st.session_state.current_mode = "Research Analysis"
            st.rerun()

    with col2:
        st.markdown(
            '<div class="card" style="text-align:center; padding:2rem; min-height:220px; cursor:pointer;">'
            '<div style="font-size:2.2rem; margin-bottom:0.6rem;">&#128172;</div>'
            '<div style="font-family:Playfair Display,serif; font-size:1.15rem; font-weight:600;'
            ' color:#3B2F24; margin-bottom:0.4rem;">Chat with Document</div>'
            '<div style="font-size:0.82rem; color:#7B5E4B; line-height:1.5;">'
            'Upload a document and ask questions — AI retrieves relevant sections to answer</div></div>',
            unsafe_allow_html=True
        )
        if st.button("Start Chat", key="land_chat", use_container_width=True):
            st.session_state.current_mode = "Chat with Document"
            st.rerun()

    with col3:
        st.markdown(
            '<div class="card" style="text-align:center; padding:2rem; min-height:220px; cursor:pointer;">'
            '<div style="font-size:2.2rem; margin-bottom:0.6rem;">&#128203;</div>'
            '<div style="font-family:Playfair Display,serif; font-size:1.15rem; font-weight:600;'
            ' color:#3B2F24; margin-bottom:0.4rem;">Compare Documents</div>'
            '<div style="font-size:0.82rem; color:#7B5E4B; line-height:1.5;">'
            'Upload 2+ documents for structured comparison of themes, findings, and differences</div></div>',
            unsafe_allow_html=True
        )
        if st.button("Start Comparison", key="land_compare", use_container_width=True):
            st.session_state.current_mode = "Compare Documents"
            st.rerun()


# ====================================================================
# MODE: RESEARCH ANALYSIS
# ====================================================================
elif mode == "Research Analysis":
    st.markdown("""
    <div class="hero">
        <h1>AI Research Analyst</h1>
        <p>Multi-agent AI system for research analysis and insight extraction</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # Input method toggle
    input_mode = st.radio("Input Method", ["Text", "File Upload"],
                          horizontal=True, label_visibility="collapsed")

    # Input section
    text_input = ""

    if input_mode == "Text":
        text_input = st.text_area(
            "Paste your research text", height=220,
            placeholder="Paste your article, research paper, notes, or any text you want analyzed...",
            label_visibility="collapsed"
        )
        if text_input:
            word_count = len(text_input.split())
            st.markdown(
                f'<div style="text-align:right; color:#A98B76; font-size:0.8rem; opacity:0.7;">'
                f'{word_count} words</div>', unsafe_allow_html=True
            )
    else:
        uploaded_file = st.file_uploader(
            "Upload your document", type=["pdf", "docx", "xlsx"],
            label_visibility="collapsed", help="Supported formats: PDF, DOCX, XLSX"
        )
        if uploaded_file:
            if uploaded_file.name.endswith(".pdf"):
                text_input = extract_pdf(uploaded_file)
            elif uploaded_file.name.endswith(".docx"):
                text_input = extract_docx(uploaded_file)
            elif uploaded_file.name.endswith(".xlsx"):
                text_input = extract_excel(uploaded_file)

            word_count = len(text_input.split())
            st.markdown(
                f'<div class="card" style="display:flex; align-items:center; gap:1rem; padding:1rem 1.2rem;">'
                f'<div style="font-size:1.5rem;">&#128196;</div>'
                f'<div><div style="color:#3B2F24; font-weight:600;">{uploaded_file.name}</div>'
                f'<div style="color:#7B5E4B; font-size:0.8rem;">{word_count} words extracted</div></div></div>',
                unsafe_allow_html=True
            )

    st.markdown("")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        analyze_clicked = st.button("Analyze Research")

    if analyze_clicked:
        if not text_input.strip():
            st.warning("Please provide some text or upload a file to analyze.")
        else:
            # Show a prominent processing banner
            status_container = st.empty()
            status_container.markdown(
                '<div class="card" style="text-align:center; padding:2rem; border:2px solid #A98B76;">'
                '<div style="font-size:1.5rem; margin-bottom:0.5rem;">&#9881;</div>'
                '<div style="font-weight:600; color:#3B2F24; font-size:1.05rem; margin-bottom:0.3rem;">'
                'Processing Your Research</div>'
                '<div style="color:#7B5E4B; font-size:0.88rem;">'
                'AI agents are planning, analyzing, and extracting insights. This may take a minute...</div>'
                '</div>',
                unsafe_allow_html=True
            )

            try:
                response = requests.post(f"{API_BASE}/analyze", json={"text": text_input}, timeout=300)
                status_container.empty()

                if response.status_code == 200:
                    data = response.json()
                    st.session_state.last_analysis = data
                    st.session_state.last_analysis_text = text_input

                    # Auto-save to history if logged in
                    if st.session_state.logged_in:
                        try:
                            requests.post(f"{API_BASE}/history/save", json={
                                "username": st.session_state.username,
                                "input_text": text_input,
                                "plan": data["plan"],
                                "analysis": data["analysis"],
                                "insights": data["insights"],
                                "web_results": data.get("web_results", [])
                            })
                        except Exception:
                            pass
                    st.rerun()
                else:
                    st.error("Backend returned an error. Please ensure the server is running.")
            except requests.exceptions.ConnectionError:
                status_container.empty()
                st.error("Could not connect to the backend. Please start the FastAPI server.")
            except Exception as e:
                status_container.empty()
                st.error(f"An unexpected error occurred: {str(e)}")

    # Show results from session state (persists across reruns)
    if st.session_state.last_analysis:
        render_results(st.session_state.last_analysis, st.session_state.last_analysis_text, show_save=False)


# ====================================================================
# MODE: CHAT WITH DOCUMENT (RAG)
# ====================================================================
elif mode == "Chat with Document":

    # If already indexed, show compact bar instead of full input UI
    if st.session_state.rag_doc_uploaded:
        doc_label = "Pasted Text" if st.session_state.rag_doc_name.startswith("pasted_text_") else st.session_state.rag_doc_name
        st.markdown(
            f'<div style="display:flex; align-items:center; background:#FFFFFF;'
            f' border:1px solid rgba(169,139,118,0.2); border-radius:12px;'
            f' padding:0.7rem 1.2rem; margin-bottom:0.5rem;">'
            f'<span style="color:#6B7D3E; font-size:1.1rem; margin-right:0.6rem;">&#9989;</span>'
            f'<span style="color:#3B2F24; font-weight:600; font-size:0.9rem;">{doc_label}</span>'
            f'<span style="color:#7B5E4B; font-size:0.78rem; margin-left:0.5rem;">— Document loaded. Ask anything below!</span>'
            f'</div>',
            unsafe_allow_html=True
        )

    else:
        st.markdown("""
        <div class="hero">
            <h1>Chat with Your Document</h1>
            <p>Provide text or upload a document, then ask questions about it</p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    if not st.session_state.rag_doc_uploaded:
        rag_input_mode = st.radio("Input Method", ["Upload File", "Paste Text"],
                                  horizontal=True, label_visibility="collapsed", key="rag_input_mode")
    else:
        rag_input_mode = None

    if not st.session_state.rag_doc_uploaded:
        if rag_input_mode == "Paste Text":
            rag_text = st.text_area(
                "Paste text to chat with", height=180,
                placeholder="Paste your research text here...",
                label_visibility="collapsed", key="rag_text_input"
            )
            if rag_text and rag_text.strip():
                text_doc_name = "pasted_text_" + str(hash(rag_text[:50]))[:8]
                if st.button("Load This Text", use_container_width=True):
                    with st.spinner("Reading your text..."):
                        try:
                            requests.post(f"{API_BASE}/rag/reset")
                            files = {"file": ("pasted_text.txt", rag_text.encode("utf-8"), "text/plain")}
                            response = requests.post(f"{API_BASE}/rag/upload", files=files)
                            if response.status_code == 200:
                                result = response.json()
                                if result["status"] == "ok":
                                    st.session_state.rag_doc_uploaded = True
                                    st.session_state.rag_doc_name = text_doc_name
                                    st.session_state.rag_chat_history = []
                                    st.rerun()
                                else:
                                    st.error(result.get("message", "Failed to index text"))
                            else:
                                st.error("Failed to send text to backend.")
                        except requests.exceptions.ConnectionError:
                            st.error("Could not connect to the backend.")
        else:
            uploaded_doc = st.file_uploader(
                "Upload a document to chat with", type=["pdf", "docx", "xlsx"],
                label_visibility="collapsed", key="rag_uploader",
                help="Upload a PDF, DOCX, or XLSX file"
            )
            if uploaded_doc:
                if uploaded_doc.name != st.session_state.rag_doc_name:
                    with st.spinner(f"Reading {uploaded_doc.name}..."):
                        try:
                            requests.post(f"{API_BASE}/rag/reset")
                            files = {"file": (uploaded_doc.name, uploaded_doc.getvalue(), uploaded_doc.type)}
                            response = requests.post(f"{API_BASE}/rag/upload", files=files)
                            if response.status_code == 200:
                                result = response.json()
                                if result["status"] == "ok":
                                    st.session_state.rag_doc_uploaded = True
                                    st.session_state.rag_doc_name = uploaded_doc.name
                                    st.session_state.rag_chat_history = []
                                    st.rerun()
                                else:
                                    st.error(result.get("message", "Failed to index document"))
                            else:
                                st.error("Failed to upload document to backend.")
                        except requests.exceptions.ConnectionError:
                            st.error("Could not connect to the backend.")
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    if st.session_state.rag_doc_uploaded:
        for msg in st.session_state.rag_chat_history:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])
                if msg["role"] == "assistant" and msg.get("sources"):
                    with st.expander("View source chunks"):
                        for src in msg["sources"]:
                            st.markdown(
                                f'<div style="background:#FAF5ED; padding:0.8rem; border-radius:8px;'
                                f' margin-bottom:0.5rem; font-size:0.85rem; color:#5C4A3A;'
                                f' border-left: 3px solid #A98B76;">'
                                f'<div style="color:#7B5E4B; font-size:0.7rem; margin-bottom:0.3rem;">'
                                f'Chunk {src["chunk_index"]} &bull; Relevance: {src["score"]}</div>'
                                f'{src["text"][:300]}{"..." if len(src["text"]) > 300 else ""}</div>',
                                unsafe_allow_html=True
                            )

        question = st.chat_input("Ask a question about your document...")
        if question:
            st.session_state.rag_chat_history.append({"role": "user", "content": question})
            with st.chat_message("user"):
                st.write(question)

            with st.chat_message("assistant"):
                with st.spinner("Searching document..."):
                    try:
                        response = requests.post(
                            f"{API_BASE}/rag/query", json={"question": question}
                        )
                        if response.status_code == 200:
                            data = response.json()
                            st.write(data["answer"])
                            if data["sources"]:
                                with st.expander("View source chunks"):
                                    for src in data["sources"]:
                                        st.markdown(
                                            f'<div style="background:#FAF5ED; padding:0.8rem; border-radius:8px;'
                                            f' margin-bottom:0.5rem; font-size:0.85rem; color:#5C4A3A;'
                                            f' border-left: 3px solid #A98B76;">'
                                            f'<div style="color:#7B5E4B; font-size:0.7rem; margin-bottom:0.3rem;">'
                                            f'Chunk {src["chunk_index"]} &bull; Relevance: {src["score"]}</div>'
                                            f'{src["text"][:300]}{"..." if len(src["text"]) > 300 else ""}</div>',
                                            unsafe_allow_html=True
                                        )
                            st.session_state.rag_chat_history.append({
                                "role": "assistant", "content": data["answer"],
                                "sources": data["sources"]
                            })

                            # Auto-save first Q&A to history
                            if st.session_state.logged_in and len(st.session_state.rag_chat_history) == 2:
                                first_q = st.session_state.rag_chat_history[0]["content"]
                                try:
                                    requests.post(f"{API_BASE}/history/save", json={
                                        "username": st.session_state.username,
                                        "input_text": first_q,
                                        "plan": f"Chat with Document: {doc_label}",
                                        "analysis": data["answer"],
                                        "insights": "",
                                        "web_results": []
                                    })
                                except Exception:
                                    pass
                            st.rerun()
                        else:
                            st.error("Backend error.")
                    except requests.exceptions.ConnectionError:
                        st.error("Could not connect to the backend.")

        if st.session_state.rag_chat_history:
            st.markdown("")
            if st.button("Clear Chat History"):
                st.session_state.rag_chat_history = []
                st.rerun()
    else:
        if rag_input_mode == "Upload File":
            st.markdown(
                '<div class="card" style="text-align:center; padding:2.5rem;">'
                '<div style="font-size:2rem; margin-bottom:0.5rem;">&#128196;</div>'
                '<div style="color:#3B2F24; font-size:1rem; font-weight:600; margin-bottom:0.3rem;">'
                'Upload a Document to Start</div>'
                '<div style="color:#7B5E4B; font-size:0.85rem;">'
                'Upload your file and start asking questions about it.</div></div>',
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                '<div class="card" style="text-align:center; padding:2.5rem;">'
                '<div style="font-size:2rem; margin-bottom:0.5rem;">&#128221;</div>'
                '<div style="color:#3B2F24; font-size:1rem; font-weight:600; margin-bottom:0.3rem;">'
                'Paste Your Text and Click "Load This Text"</div>'
                '<div style="color:#7B5E4B; font-size:0.85rem;">'
                'Paste your text, load it, and start asking questions.</div></div>',
                unsafe_allow_html=True
            )


# ====================================================================
# MODE: COMPARE DOCUMENTS
# ====================================================================
elif mode == "Compare Documents":
    st.markdown("""
    <div class="hero">
        <h1>Compare Documents</h1>
        <p>Upload multiple documents to get a structured comparison of themes and findings</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "Upload 2 or more documents to compare", type=["pdf", "docx", "xlsx"],
        accept_multiple_files=True, label_visibility="collapsed",
        key="compare_uploader", help="Upload at least 2 documents (PDF, DOCX, or XLSX)"
    )

    if uploaded_files:
        st.markdown(
            f'<div class="card"><div class="section-header">'
            f'{len(uploaded_files)} Document(s) Uploaded</div></div>',
            unsafe_allow_html=True
        )
        for f in uploaded_files:
            st.markdown(
                f'<div style="display:flex; align-items:center; gap:0.8rem; padding:0.5rem 0;'
                f' color:#3B2F24; font-size:0.92rem;">'
                f'<span style="color:#6B7D3E;">&#128196;</span> {f.name}</div>',
                unsafe_allow_html=True
            )

        st.markdown("")

        if len(uploaded_files) < 2:
            st.warning("Please upload at least 2 documents to compare.")
        else:
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                compare_clicked = st.button("Compare Documents")

            if compare_clicked:
                with st.spinner("AI is summarizing and comparing your documents..."):
                    try:
                        files_payload = [
                            ("files", (f.name, f.getvalue(), f.type))
                            for f in uploaded_files
                        ]
                        response = requests.post(f"{API_BASE}/compare", files=files_payload)
                        if response.status_code == 200:
                            data = response.json()

                            # Auto-save to history
                            if st.session_state.logged_in:
                                doc_names_str = " vs ".join(data["doc_names"])
                                summaries_text = "\n\n".join(
                                    [f"{n}:\n{data['summaries'][n]}" for n in data["doc_names"]]
                                )
                                try:
                                    requests.post(f"{API_BASE}/history/save", json={
                                        "username": st.session_state.username,
                                        "input_text": doc_names_str,
                                        "plan": f"Compare Documents: {doc_names_str}",
                                        "analysis": summaries_text,
                                        "insights": data["comparison"],
                                        "web_results": []
                                    })
                                except Exception:
                                    pass

                            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

                            tab1, tab2 = st.tabs(["Document Summaries", "Comparison"])
                            with tab1:
                                for doc_name in data["doc_names"]:
                                    summary = data["summaries"][doc_name]
                                    st.markdown(
                                        f'<div class="card"><div class="section-header">{doc_name}</div>'
                                        f'<div class="result-content">{summary}</div></div>',
                                        unsafe_allow_html=True
                                    )
                            with tab2:
                                st.markdown(
                                    f'<div class="card"><div class="section-header">Comparative Analysis</div>'
                                    f'<div class="result-content">{data["comparison"]}</div></div>',
                                    unsafe_allow_html=True
                                )
                        else:
                            st.error("Backend returned an error.")
                    except requests.exceptions.ConnectionError:
                        st.error("Could not connect to the backend. Please start the FastAPI server.")
                    except Exception as e:
                        st.error(f"An unexpected error occurred: {str(e)}")
    else:
        st.markdown(
            '<div class="card" style="text-align:center; padding:3rem;">'
            '<div style="font-size:2.5rem; margin-bottom:0.5rem;">&#128203;</div>'
            '<div style="color:#3B2F24; font-size:1.1rem; font-weight:600; margin-bottom:0.3rem;">'
            'Upload Documents to Compare</div>'
            '<div style="color:#7B5E4B; font-size:0.9rem;">'
            'Upload 2 or more research documents to get a structured comparison<br>'
            'of themes, similarities, differences, and unique findings.</div></div>',
            unsafe_allow_html=True
        )


